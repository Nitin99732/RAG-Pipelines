from docx import Document
from langchain_core.documents import Document as LCDocument


def chunk_docx_by_headings(file_path: str) -> list[LCDocument]:

    """
    Split DOCX file using heading styles.
    """

    doc = Document(file_path)

    chunks = []

    current_heading = "Untitled"

    current_text = []

    for para in doc.paragraphs:

        if para.style.name.startswith("Heading"):

            if current_text:

                chunks.append(
                    LCDocument(
                        page_content="\n".join(current_text),
                        metadata={
                            "heading": current_heading
                        }
                    )
                )

            current_heading = para.text

            current_text = []

        else:

            current_text.append(para.text)

    return chunks